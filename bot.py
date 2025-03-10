import os
import logging
import asyncio
import uuid
from google.api_core.client_options import ClientOptions
from google.cloud import documentai
from telegram import Update, InlineKeyboardMarkup, InlineKeyboardButton, InputFile
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, filters, CallbackContext
from docx import Document
from dotenv import load_dotenv

load_dotenv()

BOT_TOKEN = os.getenv("BOT_TOKEN")
PROJECT_ID = os.getenv("PROJECT_ID")
LOCATION = os.getenv("LOCATION")
PROCESSOR_ID = os.getenv("PROCESSOR_ID")
GOOGLE_CREDENTIALS = os.getenv("GOOGLE_CREDENTIALS")
SUMMARIZER_PROCESSOR_ID = os.getenv("SUMMARIZER_PROCESSOR_ID")


os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = GOOGLE_CREDENTIALS

# Create directories if they don't exist
os.makedirs("downloads", exist_ok=True)
os.makedirs("outputs", exist_ok=True)

# Logging Configuration
logging.basicConfig(format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO)


async def start(update: Update, context: CallbackContext) -> None:
    """Start Command"""
    await update.message.reply_text("👋 Send me an image or a PDF, and I'll extract the text! You can choose the output format.")


def process_document(file_path: str) -> str:
    """Processes a document using Google Document AI"""
    client_options = ClientOptions(api_endpoint=f"{LOCATION}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(client_options=client_options)

    processor_path = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{PROCESSOR_ID}"

    with open(file_path, "rb") as image:
        image_content = image.read()

    # Detect correct MIME type
    if file_path.endswith((".jpg", ".jpeg", ".png")):
        mime_type = "image/jpeg"
    elif file_path.endswith(".pdf"):
        mime_type = "application/pdf"
    else:
        return "❌ Unsupported file format."

    # Form request
    raw_document = documentai.RawDocument(content=image_content, mime_type=mime_type)
    request = documentai.ProcessRequest(name=processor_path, raw_document=raw_document)

    # Process document
    result = client.process_document(request=request)
    document = result.document
    extracted_text = document.text

    return extracted_text if extracted_text.strip() else "❌ No text recognized."


def summarize_document(file_path: str) -> str:
    """Summarizes a document using Google Document AI Summarizer"""
    client_options = ClientOptions(api_endpoint=f"{LOCATION}-documentai.googleapis.com")
    client = documentai.DocumentProcessorServiceClient(client_options=client_options)

    processor_path = f"projects/{PROJECT_ID}/locations/{LOCATION}/processors/{SUMMARIZER_PROCESSOR_ID}"

    with open(file_path, "rb") as document_file:
        document_content = document_file.read()

    # Detect correct MIME type
    if file_path.endswith((".jpg", ".jpeg", ".png")):
        mime_type = "image/jpeg"
    elif file_path.endswith(".pdf"):
        mime_type = "application/pdf"
    else:
        return "❌ Unsupported file format."

    # Form request
    raw_document = documentai.RawDocument(content=document_content, mime_type=mime_type)
    request = documentai.ProcessRequest(name=processor_path, raw_document=raw_document)

    # Process document
    result = client.process_document(request=request)
    document = result.document

    # Extract summary from the processed document
    summary = ""
    for entity in document.entities:
        if entity.type_ == "summary":
            summary += entity.mention_text + "\n"

    if not summary:
        # If no specific summary entity is found, use the text field
        summary = document.text

    return summary if summary.strip() else "❌ No summary could be generated."

def save_as_docx(text, file_name="output.docx"):
    """Saves extracted text as a DOCX file"""
    doc_path = f"outputs/{file_name}"
    doc = Document()
    doc.add_paragraph(text)
    doc.save(doc_path)
    return doc_path


async def send_text_chunks(update: Update, text: str, chunk_size=4096):
    """Splits and sends text in chunks to avoid Telegram's message limit"""
    for i in range(0, len(text), chunk_size):
        await update.effective_message.reply_text(text[i : i + chunk_size])


async def send_output_options(update: Update, context: CallbackContext, text: str):
    """Sends output format selection buttons"""
    text_key = str(uuid.uuid4())[:8]  # Create an 8-character unique key
    context.user_data[text_key] = text  # Store text in context memory

    # Store the file path in context for summarization
    if "current_file_path" in context.user_data:
        context.user_data[f"file_path_{text_key}"] = context.user_data["current_file_path"]

    keyboard = [
        [
            InlineKeyboardButton("📩 Message", callback_data=f"output_message|{text_key}"),
            InlineKeyboardButton("📄 TXT File", callback_data=f"output_txt|{text_key}"),
        ],
        [
            InlineKeyboardButton("📄+📩 TXT & Message", callback_data=f"output_both|{text_key}"),
            InlineKeyboardButton("📜 Word File (DOCX)", callback_data=f"output_docx|{text_key}"),
        ],
        [
            InlineKeyboardButton("📝 Summarize Document", callback_data=f"output_summarize|{text_key}"),
        ]
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.effective_message.reply_text("📝 Choose an output format:", reply_markup=reply_markup)

async def handle_document(update: Update, context: CallbackContext) -> None:
    """Handles document/photo processing"""
    file = update.message.document or update.message.photo[-1]  # Get file
    file_id = file.file_id
    file_name = f"downloads/{file_id}.jpg" if update.message.photo else f"downloads/{file.file_name}"

    try:
        # Download file asynchronously
        new_file = await context.bot.get_file(file_id)
        await new_file.download_to_drive(file_name)

        # Store the file path in context for later use
        context.user_data["current_file_path"] = file_name

        await update.effective_message.reply_text("🔄 Processing file, please wait...")

        # Process using Google Document AI (Synchronously for faster execution)
        loop = asyncio.get_running_loop()
        extracted_text = await loop.run_in_executor(None, process_document, file_name)

        # If no text found, notify the user
        if "❌" in extracted_text:
            await update.effective_message.reply_text(extracted_text)
            return

        # Offer output options
        await send_output_options(update, context, extracted_text)

    except Exception as e:
        logging.error(f"Error processing file: {e}")
        await update.effective_message.reply_text("❌ An error occurred while processing the document. Please try again!")


async def handle_output_choice(update: Update, context: CallbackContext) -> None:
    """Handles the output choice from inline buttons"""
    query = update.callback_query
    await query.answer()

    try:
        option, text_key = query.data.split("|")
        extracted_text = context.user_data.get(text_key, "")

        if not extracted_text:
            await query.message.reply_text("❌ Error retrieving processed text.")
            return

        txt_path = f"outputs/{text_key}.txt"
        docx_path = f"outputs/{text_key}.docx"

        # Handle summarize option
        if option == "output_summarize":
            # Get the file path from context
            file_path = context.user_data.get(f"file_path_{text_key}")

            if not file_path:
                await query.message.reply_text("❌ Cannot find the original file for summarization.")
                return

            await query.message.reply_text("🔄 Summarizing document, please wait...")

            # Process using Google Document AI Summarizer
            loop = asyncio.get_running_loop()
            summary = await loop.run_in_executor(None, summarize_document, file_path)

            if "❌" in summary:
                await query.message.reply_text(summary)
                return

            # Offer output options for the summary
            context.user_data[f"summary_{text_key}"] = summary

            summary_keyboard = [
                [
                    InlineKeyboardButton("📩 Message", callback_data=f"summary_message|{text_key}"),
                    InlineKeyboardButton("📄 TXT File", callback_data=f"summary_txt|{text_key}"),
                ],
                [
                    InlineKeyboardButton("📜 Word File (DOCX)", callback_data=f"summary_docx|{text_key}"),
                ]
            ]
            summary_markup = InlineKeyboardMarkup(summary_keyboard)
            await query.message.reply_text("📝 Choose an output format for the summary:", reply_markup=summary_markup)
            return

        # Handle summary output formats
        elif option.startswith("summary_"):
            summary = context.user_data.get(f"summary_{text_key}", "")

            if not summary:
                await query.message.reply_text("❌ Error retrieving summary.")
                return

            summary_path = f"outputs/summary_{text_key}.txt"
            summary_docx = f"outputs/summary_{text_key}.docx"

            # Save summary to TXT file
            with open(summary_path, "w", encoding="utf-8") as file:
                file.write(summary)

            if option == "summary_message":
                await send_text_chunks(update, f"📝 **Document Summary:**\n\n{summary}")

            elif option == "summary_txt":
                with open(summary_path, "rb") as doc:
                    await query.message.reply_document(InputFile(doc, filename="document_summary.txt"),
                                                       caption="📄 Here is your document summary.")

            elif option == "summary_docx":
                docx_file_path = save_as_docx(summary, file_name=f"summary_{text_key}.docx")
                with open(docx_file_path, "rb") as doc:
                    await query.message.reply_document(InputFile(doc, filename="document_summary.docx"),
                                                       caption="📜 Here is your document summary as a Word document.")
            return

        # Original options handling (unchanged)
        # Save TXT file
        with open(txt_path, "w", encoding="utf-8") as file:
            file.write(extracted_text)

        if option == "output_message":
            await send_text_chunks(update, f"📜 **Extracted Text:**\n\n{extracted_text}")

        elif option == "output_txt":
            with open(txt_path, "rb") as doc:
                await query.message.reply_document(InputFile(doc, filename="extracted_text.txt"),
                                                   caption="📄 Here is your extracted text file.")

        elif option == "output_both":
            await send_text_chunks(update, f"📜 **Extracted Text:**\n\n{extracted_text}")
            with open(txt_path, "rb") as doc:
                await query.message.reply_document(InputFile(doc, filename="extracted_text.txt"),
                                                   caption="📄 Here is your extracted text file.")

        elif option == "output_docx":
            docx_file_path = save_as_docx(extracted_text, file_name=f"{text_key}.docx")
            with open(docx_file_path, "rb") as doc:
                await query.message.reply_document(InputFile(doc, filename="extracted_text.docx"),
                                                   caption="📜 Here is your extracted text as a Word document.")

        else:
            await query.message.reply_text("❌ Invalid option.")

    except Exception as e:
        logging.error(f"Error handling output choice: {e}")
        await query.message.reply_text("❌ An error occurred while processing your selection.")

def main() -> None:
    """Runs the bot"""
    application = Application.builder().token(BOT_TOKEN).build()

    application.add_handler(CommandHandler("start", start))
    application.add_handler(MessageHandler(filters.Document.ALL | filters.PHOTO, handle_document))
    application.add_handler(CallbackQueryHandler(handle_output_choice))

    print("🤖 Bot is running!")
    application.run_polling()


if __name__ == "__main__":
    main()